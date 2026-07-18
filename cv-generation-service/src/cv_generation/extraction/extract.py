"""Safe Base CV text extraction (PDF / DOCX / Markdown)."""

from __future__ import annotations

import io
import re
import zipfile
from dataclasses import dataclass
from enum import StrEnum

from cv_generation.models.errors import ErrorCode, ServiceError


class SourceFormat(StrEnum):
    PDF = "pdf"
    DOCX = "docx"
    MARKDOWN = "markdown"
    UNKNOWN = "unknown"


@dataclass(frozen=True)
class ExtractionResult:
    text: str
    source_format: SourceFormat
    page_count: int | None = None


_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
_PHONE_RE = re.compile(
    r"(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{2,4}\)?[\s.-]?)?\d{3,4}[\s.-]?\d{3,4}"
)
_URL_RE = re.compile(r"https?://[^\s)>\]]+|www\.[^\s)>\]]+", re.IGNORECASE)

# Heuristic: scanned PDFs yield almost no text relative to page count
_MIN_CHARS_PER_PAGE = 40
_MIN_ABSOLUTE_CHARS = 30
_MAX_INPUT_PDF_PAGES = 50
_MAX_DOCX_ENTRIES = 200
_MAX_DOCX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
_MAX_DOCX_COMPRESSION_RATIO = 100.0


def detect_format(filename: str | None, content_type: str | None, data: bytes) -> SourceFormat:
    magic = _detect_magic(data)
    declared = _detect_declared(filename, content_type)

    if magic != SourceFormat.UNKNOWN and declared != SourceFormat.UNKNOWN and magic != declared:
        raise ServiceError(
            ErrorCode.MALFORMED_BASE_CV,
            "Declared Base CV format does not match file signature",
        )
    if magic != SourceFormat.UNKNOWN:
        return magic
    if declared != SourceFormat.UNKNOWN:
        return declared
    return SourceFormat.UNKNOWN


def _detect_magic(data: bytes) -> SourceFormat:
    if data[:4] == b"%PDF":
        return SourceFormat.PDF
    if data[:2] == b"PK":
        return SourceFormat.DOCX
    sample = data[:2000]
    try:
        sample.decode("utf-8")
        # Reject binary-looking payloads that happen to be mostly decodable
        if b"\x00" in sample:
            return SourceFormat.UNKNOWN
        return SourceFormat.MARKDOWN
    except UnicodeDecodeError:
        return SourceFormat.UNKNOWN


def _detect_declared(filename: str | None, content_type: str | None) -> SourceFormat:
    name = (filename or "").lower()
    ctype = (content_type or "").lower()

    if name.endswith(".md") or name.endswith(".markdown") or "markdown" in ctype or ctype == "text/plain":
        return SourceFormat.MARKDOWN
    if name.endswith(".docx") or "wordprocessingml" in ctype or "officedocument.wordprocessingml" in ctype:
        return SourceFormat.DOCX
    if name.endswith(".pdf") or ctype == "application/pdf":
        return SourceFormat.PDF
    return SourceFormat.UNKNOWN


def extract_base_cv(
    data: bytes,
    *,
    filename: str | None = None,
    content_type: str | None = None,
    max_chars: int = 100_000,
) -> ExtractionResult:
    if not data:
        raise ServiceError(
            ErrorCode.MALFORMED_BASE_CV,
            "Base CV file is empty",
        )

    source = detect_format(filename, content_type, data)
    if source == SourceFormat.UNKNOWN:
        raise ServiceError(
            ErrorCode.MALFORMED_BASE_CV,
            "Unrecognized Base CV format; expected pdf, docx, or markdown",
        )

    try:
        if source == SourceFormat.MARKDOWN:
            result = _extract_markdown(data)
        elif source == SourceFormat.DOCX:
            result = _extract_docx(data)
        else:
            result = _extract_pdf(data)
    except ServiceError:
        raise
    except Exception as exc:  # noqa: BLE001 — boundary
        raise ServiceError(
            ErrorCode.MALFORMED_BASE_CV,
            "Failed to parse Base CV",
        ) from exc

    text = result.text.strip()
    if len(text) < _MIN_ABSOLUTE_CHARS:
        raise ServiceError(
            ErrorCode.BASE_CV_NOT_EXTRACTABLE,
            "Base CV has insufficient extractable text (possibly scanned or image-only)",
        )

    if len(text) > max_chars:
        raise ServiceError(
            ErrorCode.DOCUMENT_TOO_LONG,
            f"Extracted text exceeds {max_chars} characters",
        )

    return ExtractionResult(text=text, source_format=source, page_count=result.page_count)


def _extract_markdown(data: bytes) -> ExtractionResult:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1")
    return ExtractionResult(text=text, source_format=SourceFormat.MARKDOWN)


def _assert_safe_docx_zip(data: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            infos = archive.infolist()
            if len(infos) > _MAX_DOCX_ENTRIES:
                raise ServiceError(
                    ErrorCode.MALFORMED_BASE_CV,
                    "DOCX archive has too many entries",
                )
            total_uncompressed = 0
            for info in infos:
                if info.file_size < 0 or info.compress_size < 0:
                    raise ServiceError(ErrorCode.MALFORMED_BASE_CV, "DOCX archive entry is invalid")
                if info.file_size > _MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise ServiceError(
                        ErrorCode.MALFORMED_BASE_CV,
                        "DOCX archive entry exceeds size limit",
                    )
                if info.compress_size > 0:
                    ratio = info.file_size / max(info.compress_size, 1)
                    if ratio > _MAX_DOCX_COMPRESSION_RATIO and info.file_size > 1024 * 1024:
                        raise ServiceError(
                            ErrorCode.MALFORMED_BASE_CV,
                            "DOCX archive compression ratio is unsafe",
                        )
                total_uncompressed += info.file_size
                if total_uncompressed > _MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise ServiceError(
                        ErrorCode.MALFORMED_BASE_CV,
                        "DOCX archive uncompressed size exceeds limit",
                    )
    except zipfile.BadZipFile as exc:
        raise ServiceError(ErrorCode.MALFORMED_BASE_CV, "DOCX archive is malformed") from exc


def _extract_docx(data: bytes) -> ExtractionResult:
    _assert_safe_docx_zip(data)
    from docx import Document

    document = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return ExtractionResult(text="\n".join(parts), source_format=SourceFormat.DOCX)


def _extract_pdf(data: bytes) -> ExtractionResult:
    text = ""
    page_count = 0

    # Prefer pdfplumber for layout-aware extraction; fall back to pypdf
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(data)) as pdf:
            page_count = len(pdf.pages)
            if page_count > _MAX_INPUT_PDF_PAGES:
                raise ServiceError(
                    ErrorCode.DOCUMENT_TOO_LONG,
                    f"Base CV PDF exceeds {_MAX_INPUT_PDF_PAGES} pages",
                )
            chunks: list[str] = []
            for page in pdf.pages[:_MAX_INPUT_PDF_PAGES]:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    chunks.append(page_text)
            text = "\n".join(chunks)
    except ServiceError:
        raise
    except Exception:  # noqa: BLE001
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        page_count = len(reader.pages)
        if page_count > _MAX_INPUT_PDF_PAGES:
            raise ServiceError(
                ErrorCode.DOCUMENT_TOO_LONG,
                f"Base CV PDF exceeds {_MAX_INPUT_PDF_PAGES} pages",
            )
        chunks = []
        for page in reader.pages[:_MAX_INPUT_PDF_PAGES]:
            page_text = page.extract_text() or ""
            if page_text.strip():
                chunks.append(page_text)
        text = "\n".join(chunks)

    # Scanned / image-only heuristic
    if page_count > 0 and len(text.strip()) < max(_MIN_ABSOLUTE_CHARS, page_count * _MIN_CHARS_PER_PAGE):
        raise ServiceError(
            ErrorCode.BASE_CV_NOT_EXTRACTABLE,
            "PDF appears scanned or image-only; insufficient extractable text",
        )

    return ExtractionResult(text=text, source_format=SourceFormat.PDF, page_count=page_count)


def find_emails(text: str) -> list[str]:
    return list(dict.fromkeys(_EMAIL_RE.findall(text)))


def find_phones(text: str) -> list[str]:
    candidates = _PHONE_RE.findall(text)
    # Filter short numeric noise
    cleaned = []
    for c in candidates:
        digits = re.sub(r"\D", "", c)
        if len(digits) >= 7:
            cleaned.append(c.strip())
    return list(dict.fromkeys(cleaned))


def find_urls(text: str) -> list[str]:
    return list(dict.fromkeys(_URL_RE.findall(text)))
