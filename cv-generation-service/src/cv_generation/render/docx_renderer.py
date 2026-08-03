"""DOCX renderer — single-column ATS-safe template."""

from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from cv_generation.models.canonical_cv import CanonicalCV, EducationItem, ExperienceItem

# Public ATS presentation constants (also used by renderer tests).
ATS_BODY_FONT = "Arial"
ATS_BODY_SIZE = Pt(11)
ATS_MARGIN = Inches(1)
# Usable content width on US Letter with 1" margins (8.5 - 1 - 1).
ATS_RIGHT_TAB = Inches(6.5)
# Compact spacing so a grounded ATS CV fits on one page.
_SPACE_AFTER = Pt(2)
_HEADING_SPACE_BEFORE = Pt(8)
_BULLET_LEFT = Inches(0.2)
_BULLET_HANGING = Inches(0.15)


def render_docx(cv: CanonicalCV) -> bytes:
    doc = Document()
    _apply_page_and_font_defaults(doc)

    _add_centered_header(doc, cv)

    # Always emit Professional Summary so the ATS core contract stays visible on page 1.
    summary = (cv.professional_summary or "").strip()
    heading = _heading(doc, "Professional Summary")
    heading.paragraph_format.keep_with_next = True
    summary_para = doc.add_paragraph(summary if summary else "")
    _tighten(summary_para)
    if summary_para.runs:
        _apply_ats_run(summary_para.runs[0])
    elif summary == "":
        # Empty run keeps the section present without inventing summary copy.
        _apply_ats_run(summary_para.add_run(""))

    if cv.experience:
        _heading(doc, "Experience")
        for exp in cv.experience:
            _add_experience_header(doc, exp)
            if exp.location:
                _add_body_paragraph(doc, exp.location)
            for bullet in exp.bullets:
                _add_bullet(doc, bullet)

    if cv.education:
        _heading(doc, "Education")
        for edu in cv.education:
            _add_education_block(doc, edu)

    if cv.skills:
        _heading(doc, "Skills")
        _add_body_paragraph(doc, ", ".join(cv.skills))

    if cv.awards:
        _heading(doc, "Awards/Volunteer")
        for award in cv.awards:
            _add_dated_line(doc, label=award.title, dates=award.date or "")

    if cv.projects:
        _heading(doc, "Projects")
        for proj in cv.projects:
            p = doc.add_paragraph()
            _tighten(p)
            _apply_ats_run(p.add_run(proj.name), bold=True)
            if proj.description:
                _add_body_paragraph(doc, proj.description)
            for bullet in proj.bullets:
                _add_bullet(doc, bullet)

    if cv.certifications:
        _heading(doc, "Certifications")
        for cert in cv.certifications:
            _add_bullet(doc, cert)

    if cv.languages:
        _heading(doc, "Languages")
        _add_body_paragraph(doc, ", ".join(cv.languages))

    if cv.values_alignment:
        _heading(doc, "Values Alignment")
        for item in cv.values_alignment:
            p = doc.add_paragraph()
            _tighten(p)
            _apply_ats_run(p.add_run(item.value), bold=True)
            _apply_ats_run(p.add_run(f" — {item.behaviour}"))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _apply_page_and_font_defaults(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = ATS_MARGIN
        section.bottom_margin = ATS_MARGIN
        section.left_margin = ATS_MARGIN
        section.right_margin = ATS_MARGIN

    style = doc.styles["Normal"]
    style.font.name = ATS_BODY_FONT
    style.font.size = ATS_BODY_SIZE
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ATS_BODY_FONT)
    rfonts.set(qn("w:hAnsi"), ATS_BODY_FONT)
    rfonts.set(qn("w:eastAsia"), ATS_BODY_FONT)
    rfonts.set(qn("w:cs"), ATS_BODY_FONT)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = _SPACE_AFTER
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    # Exact single-spacing in twips (12pt) keeps density predictable across viewers.
    pf.line_spacing = Twips(240)


def _add_centered_header(doc: Document, cv: CanonicalCV) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tighten(title, space_after=Pt(0))
    title.paragraph_format.keep_with_next = True
    _apply_ats_run(title.add_run(cv.full_name), bold=True)

    c = cv.contact
    contact_bits = [v for v in (c.phone, c.email, c.location, c.linkedin, c.github) if v]
    if contact_bits:
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _tighten(contact, space_after=Pt(0))
        contact.paragraph_format.keep_with_next = True
        _apply_ats_run(contact.add_run(" | ".join(contact_bits)))

    if c.portfolio:
        portfolio = doc.add_paragraph()
        portfolio.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _tighten(portfolio, space_after=Pt(2))
        portfolio.paragraph_format.keep_with_next = True
        _apply_ats_run(portfolio.add_run(c.portfolio))


def _add_experience_header(doc: Document, exp: ExperienceItem) -> None:
    label = f"{exp.title}, {exp.company}"
    dates = " - ".join(filter(None, [exp.start_date, exp.end_date]))
    _add_dated_line(doc, label=label, dates=dates)


def _add_education_block(doc: Document, edu: EducationItem) -> None:
    degree_bits = [bit for bit in (edu.degree, edu.field) if bit]
    label = ", ".join(degree_bits) if degree_bits else edu.institution
    # ATS template uses Month/Year of Completion — never a start-date fallback.
    completion = edu.end_date or ""
    _add_dated_line(doc, label=label, dates=completion)
    if degree_bits and edu.institution:
        _add_body_paragraph(doc, edu.institution)
    for detail in edu.details:
        _add_bullet(doc, detail)


def _add_dated_line(doc: Document, *, label: str, dates: str) -> None:
    p = doc.add_paragraph()
    _tighten(p)
    p.paragraph_format.tab_stops.add_tab_stop(ATS_RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
    _apply_ats_run(p.add_run(label), bold=True, italic=True)
    if dates:
        p.add_run("\t")
        _apply_ats_run(p.add_run(dates), bold=True, italic=True)


def _heading(doc: Document, text: str) -> Paragraph:
    p = doc.add_paragraph()
    _tighten(p, space_before=_HEADING_SPACE_BEFORE, space_after=_SPACE_AFTER)
    _apply_ats_run(p.add_run(text.upper()), bold=True)
    return p


def _add_body_paragraph(doc: Document, text: str) -> Paragraph:
    p = doc.add_paragraph()
    _tighten(p)
    _apply_ats_run(p.add_run(text))
    return p


def _add_bullet(doc: Document, text: str) -> Paragraph:
    """Compact bullet without Word List Bullet style (avoids bulky list indents)."""
    p = doc.add_paragraph()
    _tighten(p)
    p.paragraph_format.left_indent = _BULLET_LEFT
    p.paragraph_format.first_line_indent = -_BULLET_HANGING
    _apply_ats_run(p.add_run(f"• {text}"))
    return p


def _tighten(
    paragraph: Paragraph,
    *,
    space_before: Pt | None = Pt(0),
    space_after: Pt | None = _SPACE_AFTER,
) -> None:
    pf = paragraph.paragraph_format
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Twips(240)


def _apply_ats_run(run: Run, *, bold: bool = False, italic: bool = False) -> None:
    """Pin run typography to the ATS template (Arial 11pt)."""
    run.bold = bold
    run.italic = italic
    run.font.name = ATS_BODY_FONT
    run.font.size = ATS_BODY_SIZE
