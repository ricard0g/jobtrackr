"""Renderer golden-ish tests."""

from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt

from cv_generation.models.canonical_cv import (
    AwardItem,
    CanonicalCV,
    ContactInfo,
    EducationItem,
    ProjectItem,
    ValuesAlignmentItem,
)
from cv_generation.render.docx_renderer import (
    ATS_BODY_FONT,
    ATS_BODY_SIZE,
    ATS_MARGIN,
    ATS_RIGHT_TAB,
    render_docx,
)
from cv_generation.render.markdown import render_markdown
from cv_generation.render.pdf_renderer import _to_html, render_pdf
from fixtures.ats_canonical_cv import ats_core_canonical_cv, ats_trailing_canonical_cv

_CORE_HEADINGS = (
    "PROFESSIONAL SUMMARY",
    "EXPERIENCE",
    "EDUCATION",
    "SKILLS",
)
_TRAILING_HEADINGS = (
    "AWARDS/VOLUNTEER",
    "PROJECTS",
    "CERTIFICATIONS",
    "LANGUAGES",
    "VALUES ALIGNMENT",
)
_ATS_HEADINGS = _CORE_HEADINGS + _TRAILING_HEADINGS


def _docx_from_cv(cv: CanonicalCV | None = None) -> Document:
    return Document(io.BytesIO(render_docx(cv or ats_core_canonical_cv())))


def _paragraph_starting_with(doc: Document, prefix: str):
    for para in doc.paragraphs:
        if para.text.startswith(prefix):
            return para
    raise AssertionError(f"no paragraph starting with {prefix!r}")


def _markdown_h2_headings(text: str) -> list[str]:
    return [
        line[3:].strip()
        for line in text.splitlines()
        if line.startswith("## ")
    ]


def _pdf_text(cv: CanonicalCV) -> str:
    from pypdf import PdfReader

    return "\n".join(
        (page.extract_text() or "")
        for page in PdfReader(io.BytesIO(render_pdf(cv))).pages
    )


def _heading_order_in_text(text: str) -> list[str]:
    found: list[str] = []
    for heading in _ATS_HEADINGS:
        idx = text.find(heading)
        if idx >= 0:
            found.append((idx, heading))
    found.sort(key=lambda item: item[0])
    return [heading for _, heading in found]


def test_markdown_contains_name_and_email():
    data = render_markdown(ats_core_canonical_cv())
    text = data.decode("utf-8")
    assert "Ada Lovelace" in text
    assert "ada@example.com" in text
    assert "Analytical Engines" in text


def test_markdown_core_section_order_is_ats_structure():
    text = render_markdown(ats_core_canonical_cv()).decode("utf-8")
    assert _markdown_h2_headings(text) == list(_CORE_HEADINGS)
    assert "Software Engineer, Analytical Engines" in text
    assert "2020-01 - Present" in text
    assert "BSc, Mathematics" in text
    assert "2019-06" in text
    assert "2016-09" not in text
    assert "University of London" in text


def test_markdown_trailing_section_order_when_all_present():
    text = render_markdown(ats_trailing_canonical_cv()).decode("utf-8")
    headings = _markdown_h2_headings(text)
    trailing = [h for h in headings if h in _TRAILING_HEADINGS]
    assert trailing == list(_TRAILING_HEADINGS)
    assert headings.index("SKILLS") < headings.index("AWARDS/VOLUNTEER")
    assert "Ada Lovelace Award" in text and "2022" in text
    assert "Curiosity" in text and "Documented novel algorithms" in text


def test_markdown_omits_empty_trailing_sections_but_keeps_relative_order():
    cv = ats_core_canonical_cv()
    cv.projects = [ProjectItem(name="Difference Engine")]
    cv.languages = ["English"]
    headings = _markdown_h2_headings(render_markdown(cv).decode("utf-8"))
    present = [h for h in headings if h in _TRAILING_HEADINGS]
    assert present == ["PROJECTS", "LANGUAGES"]
    assert "AWARDS/VOLUNTEER" not in headings
    assert "CERTIFICATIONS" not in headings
    assert "VALUES ALIGNMENT" not in headings
    assert headings.index("SKILLS") < headings.index("PROJECTS") < headings.index(
        "LANGUAGES"
    )


def test_pdf_core_section_order_is_ats_structure():
    text = _pdf_text(ats_core_canonical_cv())
    assert _heading_order_in_text(text) == list(_CORE_HEADINGS)
    assert text.index("Ada Lovelace") < text.index("PROFESSIONAL SUMMARY")
    assert "Software Engineer, Analytical Engines" in text
    assert "2020-01 - Present" in text
    assert "BSc, Mathematics" in text
    assert "2019-06" in text
    assert "2016-09" not in text
    assert "University of London" in text


def test_pdf_trailing_section_order_when_all_present():
    text = _pdf_text(ats_trailing_canonical_cv())
    assert _heading_order_in_text(text) == list(_ATS_HEADINGS)
    assert "Ada Lovelace Award" in text and "2022" in text
    assert "Curiosity" in text and "Documented novel algorithms" in text


def test_pdf_omits_empty_trailing_sections_but_keeps_relative_order():
    cv = ats_core_canonical_cv()
    cv.projects = [ProjectItem(name="Difference Engine")]
    cv.languages = ["English"]
    text = _pdf_text(cv)
    assert _heading_order_in_text(text) == [
        *_CORE_HEADINGS,
        "PROJECTS",
        "LANGUAGES",
    ]


def test_pdf_style_signals_approximate_ats_template():
    html = _to_html(ats_core_canonical_cv())
    assert "text-align: center" in html
    assert "text-transform: uppercase" in html
    assert "font-family: Arial" in html
    assert "font-size: 11pt" in html
    assert "font-weight: bold" in html
    assert "font-style: italic" in html
    assert "<h1>Ada Lovelace</h1>" in html
    assert "<p class='contact'>" in html
    assert "class='dates'" in html
    assert html.index("<h1>") < html.index("<h2>Professional Summary</h2>")
    assert html.index("<h2>Experience</h2>") < html.index("<h2>Education</h2>")
    assert html.index("<h2>Education</h2>") < html.index("<h2>Skills</h2>")
    # Experience / education hierarchy: bold+italic job line via h3 + floated dates.
    assert "<h3>Software Engineer, Analytical Engines<span class='dates'>" in html
    assert "<h3>BSc, Mathematics<span class='dates'>2019-06</span></h3>" in html


def test_formats_share_ats_section_order_and_selected_content():
    cv = ats_trailing_canonical_cv()
    docx_text = "\n".join(p.text for p in _docx_from_cv(cv).paragraphs)
    md_text = render_markdown(cv).decode("utf-8")
    pdf_text = _pdf_text(cv)

    docx_headings = [t.strip() for t in docx_text.splitlines() if t.strip() in _ATS_HEADINGS]
    assert docx_headings == list(_ATS_HEADINGS)
    assert _markdown_h2_headings(md_text) == list(_ATS_HEADINGS)
    assert _heading_order_in_text(pdf_text) == list(_ATS_HEADINGS)

    for blob in (docx_text, md_text, pdf_text):
        assert "Ada Lovelace" in blob
        assert "ada@example.com" in blob
        assert "Software engineer with Python experience." in blob
        assert "Software Engineer, Analytical Engines" in blob
        assert "Python, FastAPI, PostgreSQL" in blob
        assert "Difference Engine" in blob
        assert "AWS Certified Developer" in blob
        assert "Curiosity" in blob


def test_formats_share_education_details_when_present():
    cv = ats_core_canonical_cv()
    cv.education[0].details = ["First-class honours"]
    docx_text = "\n".join(p.text for p in _docx_from_cv(cv).paragraphs)
    md_text = render_markdown(cv).decode("utf-8")
    pdf_text = _pdf_text(cv)
    for blob in (docx_text, md_text, pdf_text):
        assert "First-class honours" in blob


def test_docx_reopen_finds_name_and_email():
    doc = _docx_from_cv()
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Ada Lovelace" in text
    assert "ada@example.com" in text
    assert "Python" in text


def test_docx_core_section_order_is_ats_structure():
    doc = _docx_from_cv()
    headings = {
        p.text.strip(): p
        for p in doc.paragraphs
        if p.text.strip() in _CORE_HEADINGS
    }
    assert list(headings) == list(_CORE_HEADINGS)
    for para in headings.values():
        run = para.runs[0]
        assert run.bold is True
        assert run.font.name == ATS_BODY_FONT
        assert run.font.size == ATS_BODY_SIZE


def test_docx_header_is_centered_with_optional_portfolio_line():
    doc = _docx_from_cv()
    paragraphs = doc.paragraphs
    assert paragraphs[0].text.strip() == "Ada Lovelace"
    assert paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    name_run = paragraphs[0].runs[0]
    assert name_run.bold is True
    assert name_run.font.name == ATS_BODY_FONT
    assert name_run.font.size == ATS_BODY_SIZE

    contact = paragraphs[1]
    assert contact.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert "+1-555-0100" in contact.text
    assert "ada@example.com" in contact.text
    assert "London, UK" in contact.text
    assert "https://ada.dev" not in contact.text

    portfolio = paragraphs[2]
    assert portfolio.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert portfolio.text.strip() == "https://ada.dev"


def test_docx_experience_and_education_use_bold_italic_right_tab_dates():
    doc = _docx_from_cv()
    experience = _paragraph_starting_with(doc, "Software Engineer, Analytical Engines")
    education = _paragraph_starting_with(doc, "BSc, Mathematics")

    assert "2020-01 - Present" in experience.text
    assert "\t" in experience.text
    assert "2019-06" in education.text
    assert "\t" in education.text
    assert "2016-09" not in education.text

    for para in (experience, education):
        tabs = list(para.paragraph_format.tab_stops)
        assert len(tabs) == 1
        assert tabs[0].alignment == WD_TAB_ALIGNMENT.RIGHT
        assert tabs[0].position == ATS_RIGHT_TAB
        content_runs = [r for r in para.runs if r.text and r.text != "\t"]
        assert content_runs
        assert all(r.bold and r.italic for r in content_runs)
        assert all(r.font.name == ATS_BODY_FONT for r in content_runs)
        assert all(r.font.size == ATS_BODY_SIZE for r in content_runs)

    assert "University of London" in [p.text for p in doc.paragraphs]


def test_docx_education_uses_completion_date_only():
    cv = ats_core_canonical_cv()
    cv.education = [
        EducationItem(
            institution="University of London",
            degree="BSc",
            field="Mathematics",
            start_date="2016-09",
            end_date=None,
        )
    ]
    doc = _docx_from_cv(cv)
    education = _paragraph_starting_with(doc, "BSc, Mathematics")
    assert "2016-09" not in education.text
    assert "\t" not in education.text


def test_docx_uses_arial_one_inch_margins_and_no_tables_or_chrome():
    doc = _docx_from_cv()
    normal = doc.styles["Normal"]
    assert normal.font.name == ATS_BODY_FONT
    assert normal.font.size == ATS_BODY_SIZE
    assert ATS_BODY_SIZE == Pt(11)
    assert ATS_MARGIN.inches == 1.0

    assert len(doc.sections) == 1
    section = doc.sections[0]
    for margin in (
        section.left_margin,
        section.right_margin,
        section.top_margin,
        section.bottom_margin,
    ):
        assert margin == ATS_MARGIN

    cols = section._sectPr.find(qn("w:cols"))
    if cols is not None and cols.get(qn("w:num")) is not None:
        assert cols.get(qn("w:num")) == "1"

    assert doc.tables == []
    assert list(doc.inline_shapes) == []
    assert doc.part.element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}wsp") == []
    assert all(not p.text.strip() for p in section.header.paragraphs)
    assert all(not p.text.strip() for p in section.footer.paragraphs)


def test_docx_header_then_professional_summary_before_experience():
    texts = [p.text.strip() for p in _docx_from_cv().paragraphs if p.text.strip()]
    assert texts[0] == "Ada Lovelace"
    assert "PROFESSIONAL SUMMARY" in texts
    assert texts.index("PROFESSIONAL SUMMARY") < texts.index("EXPERIENCE")
    assert texts.index("PROFESSIONAL SUMMARY") <= 4


def test_docx_and_pdf_fit_on_one_page():
    from pypdf import PdfReader

    from cv_generation.models.specification import OutputFormat
    from cv_generation.render.verify import verify_rendered

    cv = ats_trailing_canonical_cv()
    docx = render_docx(cv)
    pdf = render_pdf(cv)
    assert len(PdfReader(io.BytesIO(pdf)).pages) == 1
    verify_rendered(
        pdf,
        OutputFormat.PDF,
        expected_name="Ada Lovelace",
        expected_email="ada@example.com",
    )
    verify_rendered(
        docx,
        OutputFormat.DOCX,
        expected_name="Ada Lovelace",
        expected_email="ada@example.com",
    )


def test_docx_omits_empty_core_sections_but_keeps_relative_order():
    cv = CanonicalCV(
        full_name="Ada Lovelace",
        contact=ContactInfo(email="ada@example.com"),
        professional_summary="Summary text.",
        skills=["Python"],
        experience=[],
        education=[],
        output_language="en",
    )
    texts = [p.text.strip() for p in _docx_from_cv(cv).paragraphs]
    assert "PROFESSIONAL SUMMARY" in texts
    assert "SKILLS" in texts
    assert "EXPERIENCE" not in texts
    assert "EDUCATION" not in texts
    assert texts.index("PROFESSIONAL SUMMARY") < texts.index("SKILLS")


def test_docx_trailing_section_order_when_all_present():
    doc = _docx_from_cv(ats_trailing_canonical_cv())
    texts = [p.text.strip() for p in doc.paragraphs]
    trailing = [t for t in texts if t in _TRAILING_HEADINGS]
    assert trailing == list(_TRAILING_HEADINGS)
    assert texts.index("SKILLS") < texts.index("AWARDS/VOLUNTEER")
    for heading in _TRAILING_HEADINGS:
        run = next(p.runs[0] for p in doc.paragraphs if p.text.strip() == heading)
        assert run.bold is True
        assert run.font.name == ATS_BODY_FONT
        assert run.font.size == ATS_BODY_SIZE


def test_docx_omits_empty_trailing_sections_but_keeps_relative_order():
    cv = ats_core_canonical_cv()
    cv.projects = [ProjectItem(name="Difference Engine")]
    cv.languages = ["English"]
    texts = [p.text.strip() for p in _docx_from_cv(cv).paragraphs]
    present = [t for t in texts if t in _TRAILING_HEADINGS]
    assert present == ["PROJECTS", "LANGUAGES"]
    assert "AWARDS/VOLUNTEER" not in texts
    assert "CERTIFICATIONS" not in texts
    assert "VALUES ALIGNMENT" not in texts
    assert texts.index("SKILLS") < texts.index("PROJECTS") < texts.index("LANGUAGES")


def test_docx_awards_use_item_left_date_right_when_dated():
    cv = ats_core_canonical_cv()
    cv.awards = [
        AwardItem(title="Ada Lovelace Award", date="2022"),
        AwardItem(title="Volunteer tutor, Coding Club"),
    ]
    doc = _docx_from_cv(cv)
    dated = _paragraph_starting_with(doc, "Ada Lovelace Award")
    undated = _paragraph_starting_with(doc, "Volunteer tutor, Coding Club")

    assert "2022" in dated.text
    assert "\t" in dated.text
    tabs = list(dated.paragraph_format.tab_stops)
    assert len(tabs) == 1
    assert tabs[0].alignment == WD_TAB_ALIGNMENT.RIGHT
    assert tabs[0].position == ATS_RIGHT_TAB
    content_runs = [r for r in dated.runs if r.text and r.text != "\t"]
    assert content_runs
    assert all(r.bold and r.italic for r in content_runs)

    assert "\t" not in undated.text
    assert undated.text.strip() == "Volunteer tutor, Coding Club"


def test_docx_values_alignment_renders_value_and_behaviour():
    cv = ats_core_canonical_cv()
    cv.values_alignment = [
        ValuesAlignmentItem(
            value="Curiosity",
            behaviour="Documented novel algorithms for the Analytical Engine",
        )
    ]
    texts = [p.text.strip() for p in _docx_from_cv(cv).paragraphs]
    assert "VALUES ALIGNMENT" in texts
    assert any(
        "Curiosity" in t and "Documented novel algorithms" in t for t in texts
    )
